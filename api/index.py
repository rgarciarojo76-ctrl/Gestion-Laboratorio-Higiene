import json
import os
import io
import base64
import requests
from datetime import datetime
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import urllib.parse
from docx import Document
from docx.shared import Pt
import openpyxl

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import cm
from .pdf_engine import create_premium_ficha

app = Flask(__name__)
CORS(app)

# Load configuration from environment
# In Vercel, these must be set in the project settings
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "DPTAspy2026")
GITHUB_TOKEN = os.environ.get("GITHUB_TOKEN", "")
GITHUB_OWNER = "rgarciarojo76-ctrl"
GITHUB_REPO = "Gestion-Laboratorio-Higiene"

# Path to the data inside the Vercel deployment (read-only)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(BASE_DIR)
LOCAL_DATA_PATH = os.path.join(PROJECT_DIR, "public", "contaminantes.json")
F01655_PATH = os.path.join(PROJECT_DIR, "F01655.docx")
F00662_PATH = os.path.join(PROJECT_DIR, "F00662_converted.docx")


def check_admin_password():
    """Validate admin password from header."""
    pw = request.headers.get("X-Admin-Password", "")
    return pw.strip() == ADMIN_PASSWORD.strip()


# --- Link Resolver Logic ---

def resolve_mta_url(mta_code):
    query = f'site:insst.es "{mta_code}" filetype:pdf'
    encoded_query = urllib.parse.quote(query)
    return f"https://www.google.com/search?q={encoded_query}&btnI=I"

def resolve_apa_url(search_term):
    clean_term = search_term.split('(')[0].split('->')[0].strip()
    # Use native YouTube search instead of strict Google site search
    query = f"tutorial muestreo higiene industrial {clean_term}"
    encoded_query = urllib.parse.quote_plus(query)
    return f"https://www.youtube.com/results?search_query={encoded_query}"


# --- GitHub API Integration ---

def get_file_from_github(filepath):
    """Fetch the current file content and SHA from GitHub."""
    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{filepath}"
    headers = {"Accept": "application/vnd.github.v3+json"}
    if GITHUB_TOKEN:
        headers["Authorization"] = f"token {GITHUB_TOKEN}"
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        data = response.json()
        if data.get('content'):
            content = base64.b64decode(data['content']).decode('utf-8')
        elif data.get('download_url'):
            # Fetch from raw download link (files > 1MB)
            headers_dl = {}
            if GITHUB_TOKEN:
                headers_dl["Authorization"] = f"token {GITHUB_TOKEN}"
            resp_dl = requests.get(data['download_url'], headers=headers_dl)
            content = resp_dl.text if resp_dl.status_code == 200 else ""
        else:
            content = ""
        return content, data['sha']
    return None, None


def update_file_in_github(filepath, content, commit_message, sha):
    """Commit the updated file content back to GitHub."""
    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/{filepath}"
    headers = {
        "Authorization": f"token {GITHUB_TOKEN}",
        "Accept": "application/vnd.github.v3+json",
        "Content-Type": "application/json"
    }
    encoded_content = base64.b64encode(content.encode('utf-8')).decode('utf-8')
    payload = {
        "message": commit_message,
        "content": encoded_content,
        "branch": "main"
    }
    if sha:
        payload["sha"] = sha
        
    response = requests.put(url, headers=headers, json=payload)
    return response.status_code in [200, 201]


def save_contaminants_to_github(contaminants_data, commit_message="Panel Admin: Actualizar base de datos"):
    """Saves the JSON to both data/ and public/ directories via GitHub API.
    Each file's SHA is fetched immediately before writing that file.
    data/ and public/ are in different tree nodes so their SHAs are independent.
    """
    if not GITHUB_TOKEN:
        print("ERROR: GITHUB_TOKEN no configurado en Vercel.")
        return False
        
    json_str = json.dumps(contaminants_data, indent=2, ensure_ascii=False)
    
    for path in ["data/contaminantes.json", "public/contaminantes.json"]:
        # Always fetch the SHA fresh immediately before writing this specific file
        _, sha = get_file_from_github(path)
        if not sha:
            print(f"Could not retrieve SHA for {path} — cannot write.")
            return False
        
        success = update_file_in_github(path, json_str, commit_message, sha)
        if not success:
            # Retry once with a fresh SHA in case of transient conflict
            print(f"First write to {path} failed, retrying with fresh SHA...")
            _, sha2 = get_file_from_github(path)
            if sha2:
                success = update_file_in_github(path, json_str, commit_message, sha2)
            if not success:
                print(f"Second write to {path} also failed.")
                return False
            
    return True



def log_admin_action(action_text):
    """Appends an action to the log file in GitHub."""
    if not GITHUB_TOKEN: return
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    new_line = f"[{timestamp}] {action_text}"
    
    content, sha = get_file_from_github("data/log_actividad.txt")
    # Append to existing or create new
    updated_content = ((content.strip() + "\n") if content else "") + new_line
    update_file_in_github("data/log_actividad.txt", updated_content, "Log: " + action_text, sha)


def load_local_contaminants():
    """Load contaminants. Prefer GitHub (Live) if token is available, else local filesystem."""
    if GITHUB_TOKEN:
        try:
            content, _ = get_file_from_github("public/contaminantes.json")
            if content:
                return json.loads(content)
        except Exception as e:
            print(f"GitHub fetch failed, using local: {e}")

    try:
        with open(LOCAL_DATA_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"Local load failed: {e}")
        return []


def update_excel_in_github(cas, codigo_soporte, codigo_soporte_alt, commit_message):
    """Update the Support Codes in the Excel file stored in GitHub."""
    if not GITHUB_TOKEN or not cas:
        return False
        
    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents/data/Base%20de%20conocimiento%20completa%202026.xlsx"
    headers = {
        "Accept": "application/vnd.github.v3+json",
        "Authorization": f"token {GITHUB_TOKEN}"
    }
    resp = requests.get(url, headers=headers)
    if resp.status_code != 200:
        return False
        
    data = resp.json()
    dl_url = data.get("download_url")
    if not dl_url: return False
    
    resp_dl = requests.get(dl_url, headers=headers)
    if resp_dl.status_code != 200:
        return False
        
    try:
        wb = openpyxl.load_workbook(io.BytesIO(resp_dl.content))
        ws = wb["Intranet Cargar PRUEBAS"]
        modified = False
        
        for row in ws.iter_rows(min_row=2):
            row_cas = str(row[0].value or "").strip()
            if row_cas == cas:
                curr_18 = str(row[18].value or "").strip()
                curr_19 = str(row[19].value or "").strip()
                
                new_18 = str(codigo_soporte or "").strip()
                new_19 = str(codigo_soporte_alt or "").strip()
                
                if curr_18 != new_18:
                    row[18].value = new_18 if new_18 else None
                    modified = True
                if curr_19 != new_19:
                    row[19].value = new_19 if new_19 else None
                    modified = True
                    
        if modified:
            buf = io.BytesIO()
            wb.save(buf)
            buf.seek(0)
            encoded = base64.b64encode(buf.read()).decode("utf-8")
            payload = {
                "message": commit_message,
                "content": encoded,
                "sha": data["sha"],
                "branch": "main"
            }
            put_resp = requests.put(url, headers=headers, json=payload)
            return put_resp.status_code in [200, 201]
            
    except Exception as e:
        print(f"Error modifying Excel: {e}")
        return False
        
    return True

# --- Document Generation ---

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

def _get_xml_tcs(row):
    """Get actual <tc> elements from a row, bypassing python-docx merged cell virtualization."""
    from docx.table import _Cell
    return [_Cell(tc, row._tr) for tc in row._tr.findall(f'{W}tc')]


def _set_cell_text_value(cell, text):
    """Write a value into a cell's first run, preserving font formatting.
    If the cell has placeholder en-spaces (\\u2002), clears them first."""
    if not cell.paragraphs:
        return
    # Clear all paragraphs' runs first, then write to the first run
    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = str(text)
        for extra_run in first_para.runs[1:]:
            extra_run.text = ""
    else:
        first_para.text = str(text)
    # Clear additional paragraphs (en-space placeholders)
    for para in cell.paragraphs[1:]:
        for run in para.runs:
            run.text = ""


def _set_cell_text(cell, text, preserve_format=True):
    """Set text in a cell while preserving the original font formatting."""
    if not cell.paragraphs:
        return
    paragraph = cell.paragraphs[0]
    if preserve_format and paragraph.runs:
        # Preserve the formatting of the first run
        run = paragraph.runs[0]
        run.text = str(text)
        # Clear any additional runs
        for extra_run in paragraph.runs[1:]:
            extra_run.text = ""
    else:
        paragraph.text = str(text)


def _is_category_row(row):
    """Check if a table row is a category header (all cells have the same non-empty text)."""
    cells_text = [cell.text.strip() for cell in row.cells]
    first = cells_text[0]
    return bool(first) and all(c == first for c in cells_text)


def _fill_header(doc, data):
    """Fill Table 0 (header) with client data, date, and shipping type.
    
    Table 0 XML structure per row (actual <tc> elements, NOT row.cells):
      Row 1: TC0(blank) | TC1(CÓDIGO, gridSpan=2) | TC2(NOMBRE, gridSpan=4)
      Row 2: TC0(OFIC.VENTA label) | TC1(value, gridSpan=2) | TC2(value, gridSpan=3) | TC3(Persona solicita)
      Row 3: TC0(CLIENTE label) | TC1(value, gridSpan=2) | TC2(value, gridSpan=3) | TC3(Correo)
      Row 4: TC0(REMITENTE label) | TC1(value, gridSpan=2) | TC2(value, gridSpan=3) | TC3(Teléfono)
      Row 5: TC0(COMPAÑÍA label) | TC1(value, gridSpan=2) | TC2(value, gridSpan=3) | TC3(Ref.Presupuesto)
      Row 6: TC0(Fecha label, gridSpan=2) | TC1(date value, gridSpan=3) | TC2(Firma, gridSpan=2)
      Row 7: TC0(Tipo envío label, gridSpan=2) | TC1(checkboxes, gridSpan=3) | TC2(blank, gridSpan=2)
      Row 8: TC0(Dirección label, gridSpan=4) | TC1(value, gridSpan=3)
    """
    t = doc.tables[0]

    # Row 1: CÓDIGO and NOMBRE — TC1 and TC2
    codigo = data.get("codigo", "")
    nombre = data.get("nombre", "")
    tcs1 = _get_xml_tcs(t.rows[1])
    if codigo and len(tcs1) > 1:
        _set_cell_text_value(tcs1[1], f"CÓDIGO (1)   {codigo}")
    if nombre and len(tcs1) > 2:
        _set_cell_text_value(tcs1[2], f"NOMBRE   {nombre}")

    # Row 2: OFIC. VENTA value → TC1, Persona solicita → TC3
    ofic = data.get("oficina_venta", "")
    persona = data.get("persona_solicita", "")
    tcs2 = _get_xml_tcs(t.rows[2])
    if ofic and len(tcs2) > 1:
        _set_cell_text_value(tcs2[1], ofic)
    if persona and len(tcs2) > 3:
        _set_cell_text_value(tcs2[3], f"Persona solicita:\n{persona}")

    # Row 3: CLIENTE value → TC1, Correo → TC3
    cliente = data.get("cliente_cargo", "")
    email = data.get("email", "")
    tcs3 = _get_xml_tcs(t.rows[3])
    if cliente and len(tcs3) > 1:
        _set_cell_text_value(tcs3[1], cliente)
    if email and len(tcs3) > 3:
        _set_cell_text_value(tcs3[3], f"Correo electrónico:\n{email}")

    # Row 4: REMITENTE value → TC1, Teléfono → TC3
    remitente = data.get("remitente", "")
    telefono = data.get("telefono", "")
    tcs4 = _get_xml_tcs(t.rows[4])
    if remitente and len(tcs4) > 1:
        _set_cell_text_value(tcs4[1], remitente)
    if telefono and len(tcs4) > 3:
        _set_cell_text_value(tcs4[3], f"Teléfono:\n{telefono}")

    # Row 5: COMPAÑÍA value → TC1, Ref. Presupuesto → TC3
    compania = data.get("compania", "")
    contrato_odoo = data.get("contrato_odoo", "")
    numero_pedido = data.get("numero_pedido", "")
    tcs5 = _get_xml_tcs(t.rows[5])
    if compania and len(tcs5) > 1:
        _set_cell_text_value(tcs5[1], compania)
    if (contrato_odoo or numero_pedido) and len(tcs5) > 3:
        text = "Ref. Presupuesto/ Nº pedido**:"
        if contrato_odoo:
            text += f"\nNumero de Contrato Odoo: {contrato_odoo}"
        if numero_pedido:
            text += f"\nNumero de pedido: {numero_pedido}"
        _set_cell_text_value(tcs5[3], text)

    # Row 6: Fecha de solicitud → TC1 — format DD / MM / AAAA
    fecha_raw = data.get("fecha_solicitud", "")
    if fecha_raw:
        try:
            parts = fecha_raw.split("-")
            fecha_formatted = f"{parts[2]} / {parts[1]} / {parts[0]}"
        except (IndexError, AttributeError):
            fecha_formatted = fecha_raw
        tcs6 = _get_xml_tcs(t.rows[6])
        if len(tcs6) > 1:
            _set_cell_text_value(tcs6[1], fecha_formatted)

    # Row 7: Tipo de envío — mark checkbox by inserting "x" in the right run
    # TC1 has multiple runs: checkbox squares (7pt) and text runs (6pt)
    # We find the run containing "Envío a la Oficina..." or "Envío a través de MRW"
    # and prepend "x  " to mark it, preserving font sizes
    tipo = data.get("tipo_envio", "oficina")
    tcs7 = _get_xml_tcs(t.rows[7])
    if len(tcs7) > 1:
        target_cell = tcs7[1]
        def _check_cb(p):
            from lxml import etree
            from docx.oxml.ns import qn
            for cb in p._p.xpath('.//w:checkBox'):
                if not cb.findall(qn('w:checked')):
                    cb.append(etree.Element(qn('w:checked'), {qn('w:val'): '1'}))
                    
        for para in target_cell.paragraphs:
            text_str = "".join(r.text for r in para.runs)
            if tipo == "oficina" and "Envío a la Oficina" in text_str:
                _check_cb(para)
            elif tipo == "mrw" and "Envío a través de MRW" in text_str:
                _check_cb(para)
        # Fill MRW account number
        if tipo == "mrw":
            cuenta_mrw = data.get("cuenta_mrw", "")
            if cuenta_mrw:
                for para in target_cell.paragraphs:
                    for run in para.runs:
                        if "Nº de cuenta MRW:" in run.text:
                            run.text = run.text.replace(
                                "Nº de cuenta MRW:",
                                f"Nº de cuenta MRW: {cuenta_mrw}"
                            )

    # Row 8: Dirección de envío → TC1 (only for MRW)
    if tipo == "mrw":
        direccion = data.get("direccion_envio", "")
        if direccion and len(t.rows) > 8:
            tcs8 = _get_xml_tcs(t.rows[8])
            if len(tcs8) > 1:
                _set_cell_text_value(tcs8[1], direccion)


def _process_material_table(table, requested_materials):
    """
    Process a material grid table (Tables 1 or 2):
    1. Insert quantities for requested CEF codes
    2. Mark rows for purging (no requested materials)
    3. Remove empty category headers

    Args:
        table: A python-docx Table object with 8 columns (2x4 grid)
        requested_materials: dict mapping CEF code -> quantity

    Returns:
        List of row indices to delete (in reverse order for safe deletion)
    """
    rows_to_delete = []
    category_row_indices = []
    material_row_indices = []

    # First pass: identify categories and insert quantities
    for i, row in enumerate(table.rows):
        if i == 0:
            # Header row (Nº unidades | CEF | Descripción | Obs | ...)
            continue

        if _is_category_row(row):
            category_row_indices.append(i)
            continue

        # Material row — check left side (col 1) and right side (col 5)
        cells = row.cells
        left_cef = cells[1].text.strip() if len(cells) > 1 else ""
        right_cef = cells[5].text.strip() if len(cells) > 5 else ""

        left_matched = False
        right_matched = False

        if left_cef and left_cef in requested_materials:
            qty = requested_materials[left_cef]
            _set_cell_text(cells[0], str(qty))
            left_matched = True

        if right_cef and right_cef in requested_materials:
            qty = requested_materials[right_cef]
            _set_cell_text(cells[4], str(qty))
            right_matched = True

        if not left_matched and not right_matched:
            # Neither side has a requested material — mark for deletion
            rows_to_delete.append(i)
        else:
            # At least one side matched — clear the non-matched side
            if not left_matched and left_cef:
                _set_cell_text(cells[0], "")
                _set_cell_text(cells[1], "")
                _set_cell_text(cells[2], "")
                _set_cell_text(cells[3], "")
            if not right_matched and right_cef:
                _set_cell_text(cells[4], "")
                _set_cell_text(cells[5], "")
                _set_cell_text(cells[6], "")
                _set_cell_text(cells[7], "")

        material_row_indices.append(i)

    # Second pass: check if categories have any surviving materials below them
    for cat_idx in category_row_indices:
        # Find the next category or end of table
        next_cat = None
        for other_cat in category_row_indices:
            if other_cat > cat_idx:
                next_cat = other_cat
                break

        # Check if any material rows between this category and the next survive
        has_materials = False
        for mat_idx in material_row_indices:
            if mat_idx > cat_idx and (next_cat is None or mat_idx < next_cat):
                if mat_idx not in rows_to_delete:
                    has_materials = True
                    break

        if not has_materials:
            rows_to_delete.append(cat_idx)

    return sorted(rows_to_delete, reverse=True)


def _delete_table_rows(table, row_indices):
    """Delete rows from a table by their indices (must be in reverse order)."""
    tbl = table._tbl
    for idx in row_indices:
        row_element = table.rows[idx]._tr
        tbl.remove(row_element)


def _compact_between_tables(doc):
    """Remove page breaks and empty paragraphs between material tables
    to compact the document after row purging.
    
    Document body order:
      [0] P(IMPORTANTE) [1] Table0 [2-3] empty [4] P(CONSULTAR) [5] empty
      [6] Table1 [7] P(PAGE BREAK) [8] empty [9] Table2
      [10-13] empties [14] Table3 ...
    
    We remove elements [7] and [8] (page break + empty between Tables 1-2)
    and the empty paragraphs [10-13] between Tables 2-3.
    """
    body = doc.element.body
    
    # Find the positions of Table 1 and Table 2 in the body
    table1_elem = doc.tables[1]._tbl if len(doc.tables) > 1 else None
    table2_elem = doc.tables[2]._tbl if len(doc.tables) > 2 else None
    table3_elem = doc.tables[3]._tbl if len(doc.tables) > 3 else None
    
    if table1_elem is None or table2_elem is None:
        return
    
    # Collect paragraphs between Table 1 and Table 2 to remove
    elements_to_remove = []
    found_table1 = False
    
    for elem in list(body):
        if elem is table1_elem:
            found_table1 = True
            continue
        if elem is table2_elem:
            break
        if found_table1 and elem.tag.endswith('}p'):
            elements_to_remove.append(elem)
    
    # Also collect empty paragraphs between Table 2 and Table 3
    if table3_elem is not None:
        found_table2 = False
        for elem in list(body):
            if elem is table2_elem:
                found_table2 = True
                continue
            if elem is table3_elem:
                break
            if found_table2 and elem.tag.endswith('}p'):
                # Only remove empty paragraphs, not content paragraphs
                text = ''.join(t.text or '' for t in elem.iter(f'{W}t'))
                if not text.strip():
                    elements_to_remove.append(elem)
    
    for elem in elements_to_remove:
        body.remove(elem)


def generate_docx(template_path, data, is_f01655=True):
    """Fill out the F01655 DOCX template with session data and return as BytesIO buffer.

    Strategy:
    1. Fill header fields (Table 0) with client data, date, shipping type
    2. Insert requested quantities into material grid (Tables 1-2)
    3. Purge unused material rows and empty category headers
    4. Preserve Table 3 (observations) and Table 4 (reference catalog)
    """
    doc = Document(template_path)

    if is_f01655:
        # Step 1: Fill header
        _fill_header(doc, data)

        # Step 2-3: Process materials
        materials = data.get("materials", [])
        requested = {}
        for mat in materials:
            cef = mat.get("cef", "")
            if cef:
                requested[cef] = mat.get("qty", 1)

        # Process Table 1 (main materials: AMIANTO → PREPESADOS)
        if len(doc.tables) > 1:
            rows_to_del = _process_material_table(doc.tables[1], requested)
            _delete_table_rows(doc.tables[1], rows_to_del)

        # Process Table 2 (secondary materials: TUBOS → MATERIAL)
        if len(doc.tables) > 2:
            rows_to_del = _process_material_table(doc.tables[2], requested)
            _delete_table_rows(doc.tables[2], rows_to_del)

        # Step 4: Remove page breaks and empty paragraphs between tables
        # to compact the document after row purging
        _compact_between_tables(doc)

    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

            
# --- Endpoints ---

@app.route("/api/generate-ficha", methods=["POST"])
def api_generate_ficha():
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400
    try:
        contaminante = data.get("contaminante_display", data.get("contaminante", "Desconocido"))
        buffer, filename_suggestion = create_premium_ficha(data)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename_suggestion,
            mimetype="application/pdf"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/health")
def health_check():
    return jsonify({"status": "ok", "environment": "vercel_serverless"})

@app.route("/api/contaminants/all", methods=["GET"])
def get_all_contaminants():
    return jsonify(load_local_contaminants())

@app.route("/api/link/mta/<path:mta_code>", methods=["GET"])
def link_mta(mta_code):
    return jsonify({"url": resolve_mta_url(mta_code)})

@app.route("/api/link/apa/<path:term>", methods=["GET"])
def link_apa(term):
    return jsonify({"url": resolve_apa_url(term)})

@app.route("/api/generate-f01655", methods=["POST"])
def api_generate_f01655():
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400
    try:
        buffer = generate_docx(F01655_PATH, data, is_f01655=True)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"F01655_Web_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Admin Endpoints ---

@app.route("/api/admin/auth", methods=["POST"])
def admin_auth():
    pw = request.json.get("password", "") if request.is_json else ""
    if pw.strip() == ADMIN_PASSWORD.strip():
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "Contraseña incorrecta"}), 401

@app.route("/api/admin/log-bruteforce", methods=["POST"])
def admin_log_bruteforce():
    log_admin_action("ALERTA: Intento de fuerza bruta detectado en el login de administración.")
    return jsonify({"ok": True})


@app.route("/api/admin/products", methods=["GET"])
def admin_list_products():
    if not check_admin_password(): return jsonify({"error": "No autorizado"}), 401
    q = request.args.get("q", "").strip().lower()
    contaminants = load_local_contaminants()
    if q:
        contaminants = [
            c for c in contaminants
            if q in (c.get("contaminante_display") or c.get("contaminante") or "").lower()
            or q in (c.get("cas") or "").lower()
            or q in (c.get("sinonimo") or "").lower()
        ]
    return jsonify(contaminants)


@app.route("/api/admin/products/<path:product_id>/visibility", methods=["POST"])
def admin_toggle_visibility(product_id):
    """Toggle the visible_en_app field for a product. Uses POST to avoid routing
    conflict with the PUT /products/<id> update endpoint."""
    if not check_admin_password(): return jsonify({"error": "No autorizado"}), 401
    
    try:
        # Read fresh data from GitHub (Vercel local is read-only)
        content, _ = get_file_from_github("data/contaminantes.json")
        if not content:
            return jsonify({"error": "No se pudo leer el archivo desde GitHub"}), 500
            
        contaminants = json.loads(content)
        
        for c in contaminants:
            if str(c.get("id", "")) == str(product_id):
                new_val = not bool(c.get("visible_en_app", False))
                c["visible_en_app"] = new_val
                
                commit_msg = f"Panel Admin: {'Visible' if new_val else 'Oculto'} {c.get('contaminante', str(product_id))}"
                success = save_contaminants_to_github(contaminants, commit_msg)
                if success:
                    log_admin_action(f"{'Activada' if new_val else 'Desactivada'} visibilidad de {c.get('contaminante_display') or c.get('contaminante')}")
                    return jsonify({"ok": True, "visible": new_val, "id": product_id})
                else:
                    return jsonify({"error": "Error al guardar en GitHub. Verifica el token."}), 500
                    
        return jsonify({"error": f"Producto '{product_id}' no encontrado"}), 404
        
    except Exception as e:
        print(f"Error in admin_toggle_visibility: {e}")
        return jsonify({"error": f"Error interno: {str(e)}"}), 500


@app.route("/api/admin/products/<path:product_id>", methods=["PUT"])
def admin_update_product(product_id):
    """Update all editable fields of a product by its string ID."""
    if not check_admin_password(): return jsonify({"error": "No autorizado"}), 401
    
    try:
        if not request.is_json:
            return jsonify({"error": "Se esperaba JSON en el cuerpo de la petición"}), 400
        
        updated_data = request.json
        if updated_data is None:
            return jsonify({"error": "Cuerpo JSON vacío o inválido"}), 400
        
        content, _ = get_file_from_github("data/contaminantes.json")
        if not content:
            return jsonify({"error": "No se pudo leer el archivo desde GitHub"}), 500
            
        contaminants = json.loads(content)
        found = False
        product_name = ""
        
        cas_to_update = None
        old_soporte = ""
        old_soporte_alt = ""
        new_soporte = updated_data.get("codigo_soporte", "")
        new_soporte_alt = updated_data.get("codigo_soporte_alt", "")
        
        for c in contaminants:
            if str(c.get("id", "")) == str(product_id):
                cas_to_update = c.get("cas")
                old_soporte = c.get("codigo_soporte", "")
                old_soporte_alt = c.get("codigo_soporte_alt", "")
                product_name = c.get("contaminante_display") or c.get("contaminante", str(product_id))
                c.update(updated_data)
                found = True
                break
                
        if not found:
            return jsonify({"error": f"Producto '{product_id}' no encontrado"}), 404
            
        commit_msg = f"Panel Admin: Editado {product_name}"
        success = save_contaminants_to_github(contaminants, commit_msg)
        
        if success:
            if cas_to_update and (old_soporte != new_soporte or old_soporte_alt != new_soporte_alt):
                try:
                    update_excel_in_github(
                        cas=cas_to_update, 
                        codigo_soporte=new_soporte, 
                        codigo_soporte_alt=new_soporte_alt,
                        commit_message=f"Panel Admin: Soportes Excel para {cas_to_update}"
                    )
                except Exception as e:
                    print(f"Excel update failed (non-critical): {e}")
            
            log_admin_action(f"Editado {product_name}")
            return jsonify({"ok": True})
        
        return jsonify({"error": "Error al guardar en GitHub. Verifica el token o inténtalo de nuevo."}), 500
        
    except Exception as e:
        print(f"Error in admin_update_product: {e}")
        return jsonify({"error": f"Error interno: {str(e)}"}), 500


@app.route("/api/admin/products", methods=["POST"])
def admin_create_product():
    if not check_admin_password(): return jsonify({"error": "No autorizado"}), 401
    new_product = request.json

    content, sha = get_file_from_github("data/contaminantes.json")
    if not content:
        return jsonify({"error": "No se pudo leer GitHub"}), 500

    contaminants = json.loads(content)
    # Generate a new ID
    max_id = max((c.get("id", 0) for c in contaminants), default=0)
    new_product["id"] = max_id + 1
    new_product["visible_en_app"] = True
    contaminants.append(new_product)

    msg = f"Panel Admin: Creado {new_product.get('contaminante', 'Nuevo')}"
    success = save_contaminants_to_github(contaminants, msg)
    if success:
        log_admin_action(f"Creado nuevo producto: {new_product.get('contaminante_display') or new_product.get('contaminante')}")
        return jsonify({"ok": True, "product": new_product})
    return jsonify({"error": "Error guardando en GitHub"}), 500


@app.route("/api/admin/log", methods=["GET"])
def admin_log():
    if not check_admin_password(): return jsonify({"error": "No autorizado"}), 401
    try:
        content, _ = get_file_from_github("data/log_actividad.txt")
        if content:
            lines = [l.strip() for l in content.strip().split("\n") if l.strip()]
            # We return them in order, frontend can reverse if it wants, 
            # but usually we return the list and let frontend handle it.
            return jsonify({"entries": lines})
        return jsonify({"entries": []})
    except Exception as e:
        print(f"Log read error: {e}")
        return jsonify({"entries": []})


@app.route("/api/admin/deploy", methods=["POST"])
def admin_deploy():
    if not check_admin_password(): return jsonify({"error": "No autorizado"}), 401
    # Trigger Vercel deploy hook if configured
    deploy_hook = os.environ.get("VERCEL_DEPLOY_HOOK", "")
    if deploy_hook:
        try:
            resp = requests.post(deploy_hook)
            return jsonify({"ok": True, "status": resp.status_code})
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    return jsonify({"ok": True, "message": "Deploy triggered (auto via GitHub push)"})


# For vercel
if __name__ == '__main__':
    app.run(debug=True, port=8000)
