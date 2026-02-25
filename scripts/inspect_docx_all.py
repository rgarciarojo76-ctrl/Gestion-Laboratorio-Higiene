import os
from docx import Document

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
DOCX_PATH = os.path.join(BASE_DIR, "MTMHI2025v01.docx")

def inspect_all():
    doc = Document(DOCX_PATH)
    print(f"Total tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        if not table.rows: continue
        headers = [c.text.strip().replace('\n', ' ') for c in table.rows[0].cells]
        print(f"Table {i} headers: {headers}")
        if len(table.rows) > 1:
            row1 = [c.text.strip().replace('\n', ' ') for c in table.rows[1].cells]
            print(f"  Row 1: {row1}")
            
    print("\nSearching for checkmarks in paragraphs...")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "simultáneo" in text.lower() or "✓" in text or "✔" in text or "☑" in text or "x" in text.lower():
            if "simultáneo" in text.lower() or "✓" in text or "✔" in text or "☑" in text:
                print(f"Para {i}: {text[:200]}")

if __name__ == "__main__":
    inspect_all()
