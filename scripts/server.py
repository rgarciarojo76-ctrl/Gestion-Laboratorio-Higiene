#!/usr/bin/env python3
"""
server.py — Flask API server for Gestión Laboratorio Higiene Industrial

Endpoints:
  GET  /api/contaminants?q=<query>  — Search contaminants
  POST /api/generate-f01655         — Generate F01655 (Material Request) PDF
  POST /api/generate-f00662         — Generate F00662 (Chain of Custody) PDF
"""

import json
import os
import re
import subprocess
import tempfile
from datetime import datetime
from flask import Flask, jsonify, request, send_file, redirect
from flask_cors import CORS
from link_resolver import resolve_mta_url, resolve_apa_url

# python-docx for Word manipulation
from docx import Document
from docx.shared import Pt

# reportlab for PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import cm

app = Flask(__name__)
CORS(app)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(BASE_DIR)
DATA_PATH = os.path.join(PROJECT_DIR, "data", "contaminantes.json")
PUBLIC_DATA_PATH = os.path.join(PROJECT_DIR, "public", "contaminantes.json")
LOG_PATH = os.path.join(PROJECT_DIR, "data", "log_actividad.txt")
F01655_PATH = os.path.join(PROJECT_DIR, "F01655.docx")
F00662_PATH = os.path.join(PROJECT_DIR, "F00662_converted.docx")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "output")
ADMIN_PASSWORD = "DPTAspy2026"

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Load contaminant data at startup
contaminants = []
if os.path.exists(DATA_PATH):
    with open(DATA_PATH, "r", encoding="utf-8") as f:
        contaminants = json.load(f)
    print(f"✅ Loaded {len(contaminants)} contaminants from {DATA_PATH}")


# ─── Public API ──────────────────────────────────────────────────────────────

@app.route("/api/contaminants/all", methods=["GET"])
def get_all_contaminants():
    """Return the full live contaminants list (public, no auth needed)."""
    return jsonify(contaminants)


# ─── Helper Functions ───────────────────────────────────────────────────────

def search_contaminants(query):
    """Search contaminants by name, CAS, synonym."""
    if not query or len(query) < 2:
        return []
    terms = query.lower().split()
    results = []
    for c in contaminants:
        haystack = " ".join([
            c.get("contaminante", ""),
            c.get("contaminante_display", ""),
            c.get("cas", ""),
            c.get("sinonimo", ""),
            c.get("tecnica_analitica", ""),
            c.get("metodo_analisis", ""),
            c.get("soporte_captacion", ""),
        ]).lower()
        if all(t in haystack for t in terms):
            results.append(c)
    return results[:50]


def replace_text_in_paragraphs(doc, replacements):
    """Replace text placeholders in Word document paragraphs."""
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value or ""))


def replace_text_in_tables(doc, replacements):
    """Replace text placeholders in Word document tables."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(value or ""))


def docx_to_pdf(docx_path, pdf_path):
    """
    Convert a .docx file to PDF.
    Uses textutil + cupsfilter on macOS as a fallback.
    """
    try:
        # Try LibreOffice first
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir",
             os.path.dirname(pdf_path), docx_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            # LibreOffice names the output file based on input
            generated = os.path.join(
                os.path.dirname(pdf_path),
                os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
            )
            if os.path.exists(generated) and generated != pdf_path:
                os.rename(generated, pdf_path)
            return True
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: use textutil to convert to HTML, then to PDF
    try:
        html_path = docx_path.replace(".docx", ".html")
        subprocess.run(
            ["textutil", "-convert", "html", docx_path, "-output", html_path],
            capture_output=True, text=True, timeout=15
        )
        if os.path.exists(html_path):
            # Use cupsfilter or just keep as HTML
            # For now, just rename the docx as the output
            import shutil
            shutil.copy2(docx_path, pdf_path.replace(".pdf", ".docx"))
            return True
    except Exception:
        pass

    return False


# ─── API Endpoints ──────────────────────────────────────────────────────────

@app.route("/api/contaminants", methods=["GET"])
def api_search_contaminants():
    """Search contaminants by query string."""
    query = request.args.get("q", "")
    results = search_contaminants(query)
    return jsonify(results)



# --- Intelligent DOCX Filling Helpers ---

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

def _get_xml_tcs(row):
    """Get actual <tc> elements from a row, bypassing python-docx merged cell virtualization."""
    from docx.table import _Cell
    return [_Cell(tc, row._tr) for tc in row._tr.findall(f'{W}tc')]


def _set_cell_text_value(cell, text):
    """Write a value into a cell's first run, preserving font formatting."""
    if not cell.paragraphs:
        return
    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = str(text)
        for extra_run in first_para.runs[1:]:
            extra_run.text = ""
    else:
        first_para.text = str(text)
    for para in cell.paragraphs[1:]:
        for run in para.runs:
            run.text = ""


def _set_cell_text(cell, text, preserve_format=True):
    """Set text in a cell while preserving the original font formatting."""
    if not cell.paragraphs:
        return
    paragraph = cell.paragraphs[0]
    if preserve_format and paragraph.runs:
        run = paragraph.runs[0]
        run.text = str(text)
        for extra_run in paragraph.runs[1:]:
            extra_run.text = ""
    else:
        paragraph.text = str(text)


def _is_category_row(row):
    """Check if a table row is a category header (all cells have the same non-empty text)."""
    cells_text = [cell.text.strip() for cell in row.cells]
    first = cells_text[0]
    return bool(first) and all(c == first for c in cells_text)


def _fill_header(doc, data):
    """Fill Table 0 (header) with client data using XML <tc> elements directly."""
    t = doc.tables[0]

    codigo = data.get("codigo", "")
    nombre = data.get("nombre", "")
    tcs1 = _get_xml_tcs(t.rows[1])
    if codigo and len(tcs1) > 1:
        _set_cell_text_value(tcs1[1], f"CÓDIGO (1)   {codigo}")
    if nombre and len(tcs1) > 2:
        _set_cell_text_value(tcs1[2], f"NOMBRE   {nombre}")

    ofic = data.get("oficina_venta", "")
    persona = data.get("persona_solicita", "")
    tcs2 = _get_xml_tcs(t.rows[2])
    if ofic and len(tcs2) > 1:
        _set_cell_text_value(tcs2[1], ofic)
    if persona and len(tcs2) > 3:
        _set_cell_text_value(tcs2[3], f"Persona solicita:\n{persona}")

    cliente = data.get("cliente_cargo", "")
    email = data.get("email", "")
    tcs3 = _get_xml_tcs(t.rows[3])
    if cliente and len(tcs3) > 1:
        _set_cell_text_value(tcs3[1], cliente)
    if email and len(tcs3) > 3:
        _set_cell_text_value(tcs3[3], f"Correo electrónico:\n{email}")

    remitente = data.get("remitente", "")
    telefono = data.get("telefono", "")
    tcs4 = _get_xml_tcs(t.rows[4])
    if remitente and len(tcs4) > 1:
        _set_cell_text_value(tcs4[1], remitente)
    if telefono and len(tcs4) > 3:
        _set_cell_text_value(tcs4[3], f"Teléfono:\n{telefono}")

    compania = data.get("compania", "")
    ref_pres = data.get("ref_presupuesto", "")
    tcs5 = _get_xml_tcs(t.rows[5])
    if compania and len(tcs5) > 1:
        _set_cell_text_value(tcs5[1], compania)
    if ref_pres and len(tcs5) > 3:
        _set_cell_text_value(tcs5[3], f"Ref. Presupuesto/ Nº pedido**:\n{ref_pres}")

    fecha_raw = data.get("fecha_solicitud", "")
    if fecha_raw:
        try:
            parts = fecha_raw.split("-")
            fecha_formatted = f"{parts[2]} / {parts[1]} / {parts[0]}"
        except (IndexError, AttributeError):
            fecha_formatted = fecha_raw
        tcs6 = _get_xml_tcs(t.rows[6])
        if len(tcs6) > 1:
            _set_cell_text_value(tcs6[1], fecha_formatted)

    tipo = data.get("tipo_envio", "oficina")
    tcs7 = _get_xml_tcs(t.rows[7])
    if len(tcs7) > 1:
        target_cell = tcs7[1]
        for para in target_cell.paragraphs:
            for run in para.runs:
                if tipo == "oficina" and "Envío a la Oficina" in run.text:
                    run.text = "x  " + run.text
                elif tipo == "mrw" and "Envío a través de MRW" in run.text:
                    run.text = "x  " + run.text
        if tipo == "mrw":
            cuenta_mrw = data.get("cuenta_mrw", "")
            if cuenta_mrw:
                for para in target_cell.paragraphs:
                    for run in para.runs:
                        if "Nº de cuenta MRW:" in run.text:
                            run.text = run.text.replace(
                                "Nº de cuenta MRW:",
                                f"Nº de cuenta MRW: {cuenta_mrw}"
                            )

    if tipo == "mrw":
        direccion = data.get("direccion_envio", "")
        if direccion and len(t.rows) > 8:
            tcs8 = _get_xml_tcs(t.rows[8])
            if len(tcs8) > 1:
                _set_cell_text_value(tcs8[1], direccion)


def _process_material_table(table, requested_materials):
    """Process a material grid table: insert quantities, mark rows for purging."""
    rows_to_delete = []
    category_row_indices = []
    material_row_indices = []

    for i, row in enumerate(table.rows):
        if i == 0:
            continue

        if _is_category_row(row):
            category_row_indices.append(i)
            continue

        cells = row.cells
        left_cef = cells[1].text.strip() if len(cells) > 1 else ""
        right_cef = cells[5].text.strip() if len(cells) > 5 else ""

        left_matched = False
        right_matched = False

        if left_cef and left_cef in requested_materials:
            _set_cell_text(cells[0], str(requested_materials[left_cef]))
            left_matched = True

        if right_cef and right_cef in requested_materials:
            _set_cell_text(cells[4], str(requested_materials[right_cef]))
            right_matched = True

        if not left_matched and not right_matched:
            rows_to_delete.append(i)
        else:
            if not left_matched and left_cef:
                for c in range(4):
                    _set_cell_text(cells[c], "")
            if not right_matched and right_cef:
                for c in range(4, 8):
                    _set_cell_text(cells[c], "")

        material_row_indices.append(i)

    for cat_idx in category_row_indices:
        next_cat = None
        for other_cat in category_row_indices:
            if other_cat > cat_idx:
                next_cat = other_cat
                break

        has_materials = False
        for mat_idx in material_row_indices:
            if mat_idx > cat_idx and (next_cat is None or mat_idx < next_cat):
                if mat_idx not in rows_to_delete:
                    has_materials = True
                    break

        if not has_materials:
            rows_to_delete.append(cat_idx)

    return sorted(rows_to_delete, reverse=True)


def _delete_table_rows(table, row_indices):
    """Delete rows from a table by their indices (must be in reverse order)."""
    tbl = table._tbl
    for idx in row_indices:
        row_element = table.rows[idx]._tr
        tbl.remove(row_element)


@app.route("/api/generate-f01655", methods=["POST"])
def api_generate_f01655():
    """Generate F01655 Material Request document."""
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    try:
        doc = Document(F01655_PATH)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_docx = os.path.join(OUTPUT_DIR, f"F01655_{timestamp}.docx")

        # Step 1: Fill header with client data
        _fill_header(doc, data)

        # Step 2-3: Process materials — insert quantities and purge unused rows
        materials = data.get("materials", [])
        requested = {}
        for mat in materials:
            cef = mat.get("cef", "")
            if cef:
                requested[cef] = mat.get("qty", 1)

        # Process Table 1 (AMIANTO → PREPESADOS)
        if len(doc.tables) > 1:
            rows_to_del = _process_material_table(doc.tables[1], requested)
            _delete_table_rows(doc.tables[1], rows_to_del)

        # Process Table 2 (TUBOS → MATERIAL)
        if len(doc.tables) > 2:
            rows_to_del = _process_material_table(doc.tables[2], requested)
            _delete_table_rows(doc.tables[2], rows_to_del)

        doc.save(output_docx)

        # Try to convert to PDF
        output_pdf = output_docx.replace(".docx", ".pdf")
        pdf_success = docx_to_pdf(output_docx, output_pdf)

        if pdf_success and os.path.exists(output_pdf):
            return send_file(output_pdf, as_attachment=True,
                           download_name=f"F01655_{timestamp}.pdf")
        else:
            return send_file(output_docx, as_attachment=True,
                           download_name=f"F01655_{timestamp}.docx")

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate-f00662", methods=["POST"])
def api_generate_f00662():
    """Generate F00662 Chain of Custody document."""
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    try:
        doc = Document(F00662_PATH)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_docx = os.path.join(OUTPUT_DIR, f"F00662_{timestamp}.docx")

        # Add data summary sections
        doc.add_paragraph("")
        doc.add_paragraph("═" * 60)
        p = doc.add_paragraph()
        p.add_run("DATOS CUMPLIMENTADOS AUTOMÁTICAMENTE").bold = True

        # Solicitor data
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("DATOS DEL SOLICITANTE").bold = True

        solicitor_fields = [
            ("Código", data.get("codigo", "")),
            ("Nombre", data.get("nombre", "")),
            ("Cliente (Cargo)", data.get("cliente_cargo", "")),
            ("Remitente", data.get("remitente", "")),
            ("Empresa", data.get("empresa", "")),
            ("Persona solicita", data.get("persona_solicita", "")),
            ("Email", data.get("email", "")),
            ("Teléfono", data.get("telefono", "")),
        ]

        for label, value in solicitor_fields:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(str(value))

        # Collection data
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("DATOS DE LA RECOGIDA").bold = True

        collection_fields = [
            ("Recogida por", data.get("recogida_por", "")),
            ("Fecha toma", data.get("fecha_toma", "")),
            ("Hora", data.get("hora_toma", "")),
            ("Persona toma", data.get("persona_toma", "")),
            ("Procedimiento", data.get("procedimiento", "")),
        ]

        for label, value in collection_fields:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(str(value))

        # Samples
        samples = data.get("samples", [])
        for i, sample in enumerate(samples, 1):
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.add_run(f"MUESTRA {i}").bold = True

            sample_fields = [
                ("Descripción", sample.get("descripcion", "")),
                ("Ref. Soporte", sample.get("ref_soporte", "")),
                ("Punto muestreo", sample.get("punto_muestreo", "")),
                ("Caudal ini / fin", f"{sample.get('caudal_inicial', '')} / {sample.get('caudal_final', '')} L/min"),
                ("Tiempo", f"{sample.get('tiempo_min', '')} min"),
                ("Código equipo", sample.get("codigo_equipo", "")),
                ("Tipo", sample.get("tipo_muestreo", "")),
                ("Análisis solicitado", sample.get("analisis_solicitado", "")),
                ("Observaciones", sample.get("observaciones", "")),
            ]

            for label, value in sample_fields:
                p = doc.add_paragraph()
                run = p.add_run(f"  {label}: ")
                run.bold = True
                run.font.size = Pt(10)
                p.add_run(str(value))

        doc.save(output_docx)

        # Try to convert to PDF
        output_pdf = output_docx.replace(".docx", ".pdf")
        pdf_success = docx_to_pdf(output_docx, output_pdf)

        if pdf_success and os.path.exists(output_pdf):
            return send_file(output_pdf, as_attachment=True,
                           download_name=f"F00662_{timestamp}.pdf")
        else:
            return send_file(output_docx, as_attachment=True,
                           download_name=f"F00662_{timestamp}.docx")

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate-ficha", methods=["POST"])
def api_generate_ficha():
    """Generate Ficha Técnica de Procedimiento de Muestreo PDF."""
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    try:
        contaminante = data.get("contaminante_display", data.get("contaminante", "Desconocido"))
        safe_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', contaminante)
        timestamp = datetime.now().strftime("%Y%m%d")
        
        pdf_filename = f"Ficha_Muestreo_{safe_name}_{timestamp}.pdf"
        output_pdf = os.path.join(OUTPUT_DIR, pdf_filename)
        
        doc = SimpleDocTemplate(output_pdf, pagesize=A4,
                                rightMargin=2*cm, leftMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        
        Story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor("#0369a1"),
            spaceAfter=20,
            alignment=1 # Center
        )
        
        heading_style = ParagraphStyle(
            'HeadingStyle',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor("#0f172a"),
            spaceBefore=15,
            spaceAfter=10,
            borderPadding=5,
            backColor=colors.HexColor("#f1f5f9")
        )

        normal_style = styles['Normal']
        
        # Title
        Story.append(Paragraph("FICHA TÉCNICA DE PROCEDIMIENTO DE MUESTREO - 2026", title_style))
        Story.append(Spacer(1, 10))
        
        # Helper to create styled tables
        def create_table(data_matrix):
            t = Table(data_matrix, colWidths=[5*cm, 11*cm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#f8fafc")),
                ('TEXTCOLOR', (0,0), (-1,-1), colors.HexColor("#334155")),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,-1), 10),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('INNERGRID', (0,0), (-1,-1), 0.25, colors.HexColor("#cbd5e1")),
                ('BOX', (0,0), (-1,-1), 0.25, colors.HexColor("#cbd5e1")),
            ]))
            return t

        # Block 1: Identificación
        Story.append(Paragraph("1. Identificación del Agente Químico", heading_style))
        b1_data = [
            ["Contaminante", Paragraph(contaminante, normal_style)],
            ["Nº CAS", str(data.get("cas", "—"))],
            ["Sinónimos", Paragraph(str(data.get("sinonimo", "—")), normal_style)],
        ]
        Story.append(create_table(b1_data))
        Story.append(Spacer(1, 15))

        # Block 2: Parámetros de Captación
        Story.append(Paragraph("2. Parámetros de Captación", heading_style))
        b2_data = [
            ["Soporte de Captación", Paragraph(str(data.get("soporte_captacion_display", data.get("soporte_captacion", "—"))), normal_style)],
            ["Ref. Soporte", str(data.get("ref_soporte", "—"))],
            ["Técnica Analítica", Paragraph(str(data.get("tecnica_analitica", "—")), normal_style)],
            ["Ref. Técnica", str(data.get("ref_tecnica", "—"))],
            ["Método de Análisis", Paragraph(str(data.get("metodo_analisis", "—")), normal_style)],
            ["Caudal (L/min)", str(data.get("caudal", data.get("caudal_l_min", "—")))],
            ["Volumen Recomendado (L)", str(data.get("volumen_minimo", data.get("volumen_recomendado_l", "—")))],
            ["Desorción", Paragraph(str(data.get("desorcion", "—")), normal_style)]
        ]
        Story.append(create_table(b2_data))
        Story.append(Spacer(1, 15))

        # Block 3: Criterios Analíticos y Límites
        Story.append(Paragraph("3. Criterios Analíticos y Límites Ambientales", heading_style))
        b3_data = [
            ["LOQ / LOD", f"{data.get('lq', data.get('loq', '—'))} µg / {data.get('ld', data.get('lod', '—'))} µg"],
            ["VLA-ED (mg/m³)", str(data.get("vla_ed", data.get("vla_ed_mg_m3", "—")))],
            ["VLA-EC (mg/m³)", str(data.get("vla_ec", data.get("vla_ec_mg_m3", "—")))],
            ["Frases H", Paragraph(str(data.get("frases_h", "—")), normal_style)],
            ["Notas LEP", Paragraph(str(data.get("notas_lep", "—")), normal_style)]
        ]
        Story.append(create_table(b3_data))
        Story.append(Spacer(1, 25))

        # Block 4: Observaciones Técnicas y Seguridad (Preventive Logic)
        Story.append(Paragraph("Observaciones Técnicas y de Seguridad", heading_style))
        support_raw = str(data.get("soporte_captacion_display", data.get("soporte_captacion", ""))).lower()
        method_raw = str(data.get("metodo_analisis", "")).lower()
        obs_raw = str(data.get("observaciones", "")).lower()
        
        warnings_list = []
        
        # Rule 1: Filters
        if "filtro" in support_raw or "pvc" in support_raw or "celulosa" in support_raw:
            warnings_list.append("• Asegurar el cierre hermético del portafiltros. Evitar humedad excesiva antes del pesaje.")
            
        # Rule 2: Tubes
        if "tubo" in support_raw or "carbón" in support_raw or "sílice" in support_raw:
            warnings_list.append("• Romper extremos justo antes del muestreo. Respetar la flecha de dirección de flujo. Enviar tubo blanco del mismo lote sin abrir.")
            
        # Rule 3: Light sensitive
        if "luz" in method_raw or "luz" in obs_raw or "fotosen" in obs_raw or "alum" in support_raw:
            warnings_list.append("• Proteger la muestra de la luz solar directa utilizando papel de aluminio o envases opacos.")
            
        # Rule 4: Refrigeration
        intra_transporte = str(data.get("transporte", "")).lower()
        if "refrig" in intra_transporte or "frio" in intra_transporte or "nevera" in obs_raw:
            warnings_list.append("• Mantener la muestra entre 2°C y 8°C durante el transporte al laboratorio.")

        if data.get("observaciones_concepto"):
            warnings_list.append(f"• {data.get('observaciones_concepto')}")

        if not warnings_list:
            warnings_list.append("• Seguir las condiciones generales de transporte establecidas por el laboratorio.")

        for w in warnings_list:
            Story.append(Paragraph(w, normal_style))
            Story.append(Spacer(1, 5))

        # Build PDF
        doc.build(Story)
        
        if os.path.exists(output_pdf):
            return send_file(output_pdf, as_attachment=True, download_name=pdf_filename)
        else:
            return jsonify({"error": "Failed to generate PDF file"}), 500

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
@app.route("/api/link/mta/<path:mta_code>", methods=["GET"])
def link_mta(mta_code):
    """Dynamically resolves and redirects to the exact INSST PDF."""
    target_url = resolve_mta_url(mta_code)
    if target_url:
        return redirect(target_url)
    return jsonify({"error": "No url resolved"}), 404

@app.route("/api/link/apa/<path:term>", methods=["GET"])
def link_apa(term):
    """Generates an intelligent YouTube search for APA videos."""
    target_url = resolve_apa_url(term)
    if target_url:
        return redirect(target_url)
    return jsonify({"error": "No url resolved"}), 404
# ─── Admin Helper Functions ─────────────────────────────────────────────────

def check_admin_password():
    """Validate admin password from header."""
    pw = request.headers.get("X-Admin-Password", "")
    return pw.strip() == ADMIN_PASSWORD.strip()

def save_contaminants():
    """Persist contaminants to both data/ and public/ JSON files."""
    global contaminants
    for path in [DATA_PATH, PUBLIC_DATA_PATH]:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(contaminants, f, ensure_ascii=False, indent=2)

def log_activity(action, product_name="", cas=""):
    """Append an entry to the activity log."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    product_info = product_name
    if cas:
        product_info += f" (CAS: {cas})"
    entry = f"[{timestamp}] - USUARIO: Admin - ACCIÓN: {action} - PRODUCTO: {product_info}\n"
    with open(LOG_PATH, "a", encoding="utf-8") as f:
        f.write(entry)


# ─── Admin API Endpoints ────────────────────────────────────────────────────

@app.route("/api/admin/auth", methods=["POST"])
def admin_auth():
    """Validate admin password."""
    pw = request.json.get("password", "") if request.is_json else ""
    if pw.strip() == ADMIN_PASSWORD.strip():
        return jsonify({"ok": True})
    print(f"Failed login attempt with password: '{pw}'")
    return jsonify({"ok": False, "error": "Contraseña incorrecta"}), 401


@app.route("/api/admin/products", methods=["GET"])
def admin_list_products():
    """List all products (with visibility status). Supports ?q= search."""
    if not check_admin_password():
        return jsonify({"error": "No autorizado"}), 401

    q = request.args.get("q", "").strip().lower()
    results = contaminants
    if q:
        results = [
            c for c in contaminants
            if q in (c.get("contaminante_display") or c.get("contaminante") or "").lower()
            or q in (c.get("cas") or "").lower()
            or q in (c.get("sinonimo") or "").lower()
        ]
    return jsonify(results)


@app.route("/api/admin/products/<path:product_id>", methods=["PUT"])
def admin_edit_product(product_id):
    """Edit fields of an existing product."""
    if not check_admin_password():
        return jsonify({"error": "No autorizado"}), 401

    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    for c in contaminants:
        if str(c.get("id")) == str(product_id):
            # Update only provided fields (protect id and visible_en_app)
            for key, value in data.items():
                if key not in ("id",):
                    c[key] = value
            save_contaminants()
            log_activity("Edición", c.get("contaminante_display") or c.get("contaminante", ""), c.get("cas", ""))
            return jsonify({"ok": True, "product": c})

    return jsonify({"error": "Producto no encontrado"}), 404


@app.route("/api/admin/products/<path:product_id>/visibility", methods=["PUT"])
def admin_toggle_visibility(product_id):
    """Toggle visible_en_app for a product."""
    if not check_admin_password():
        return jsonify({"error": "No autorizado"}), 401

    for c in contaminants:
        if str(c.get("id")) == str(product_id):
            new_val = not c.get("visible_en_app", False)
            c["visible_en_app"] = new_val
            save_contaminants()
            action = "Visibilidad: Activado" if new_val else "Visibilidad: Desactivado"
            log_activity(action, c.get("contaminante_display") or c.get("contaminante", ""), c.get("cas", ""))
            return jsonify({"ok": True, "visible_en_app": new_val})

    return jsonify({"error": "Producto no encontrado"}), 404


@app.route("/api/admin/products", methods=["POST"])
def admin_create_product():
    """Create a new product."""
    if not check_admin_password():
        return jsonify({"error": "No autorizado"}), 401

    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    # Generate a new unique ID (timestamp-based)
    new_id = f"MANUAL-{datetime.now().strftime('%Y%m%d%H%M%S')}"

    new_product = {
        "id": str(new_id),
        "visible_en_app": False,  # New products hidden by default
        "contaminante": data.get("contaminante", ""),
        "contaminante_display": data.get("contaminante_display", data.get("contaminante", "")),
        "cas": data.get("cas", ""),
        "sinonimo": data.get("sinonimo", ""),
        "descripcion_tecnica": data.get("descripcion_tecnica", ""),
        "soporte_captacion": data.get("soporte_captacion", ""),
        "soporte_captacion_display": data.get("soporte_captacion_display", ""),
        "tecnica_analitica": data.get("tecnica_analitica", ""),
        "metodo_analisis": data.get("metodo_analisis", ""),
        "caudal": data.get("caudal", ""),
        "volumen_minimo": data.get("volumen_minimo", ""),
        "lq": data.get("lq", ""),
        "vla_ed": data.get("vla_ed", ""),
        "vla_ec": data.get("vla_ec", ""),
        "sin_metodo_disponible": data.get("sin_metodo_disponible", False),
        "is_cmr": data.get("is_cmr", False),
        "screening_perfil": data.get("screening_perfil", ""),
        "screening_desc": data.get("screening_desc", ""),
    }

    contaminants.append(new_product)
    save_contaminants()
    log_activity("Alta", new_product.get("contaminante_display", ""), new_product.get("cas", ""))
    return jsonify({"ok": True, "product": new_product}), 201


@app.route("/api/admin/deploy", methods=["POST"])
def admin_deploy_to_vercel():
    """Commit local JSON changes to GitHub to trigger Vercel deploy."""
    if not check_admin_password():
        return jsonify({"error": "No autorizado"}), 401
    
    try:
        # Save any in-memory state explicitly
        save_contaminants()
        
        # Git commands
        subprocess.run(["git", "add", "data/contaminantes.json", "public/contaminantes.json", "data/log_actividad.txt"], cwd=PROJECT_DIR, check=True)
        # Commit might fail if no changes, we ignore error
        subprocess.run(["git", "commit", "-m", "Panel Admin: Publicar cambios de visibilidad en Vercel"], cwd=PROJECT_DIR)
        
        # Push to trigger deploy
        result = subprocess.run(["git", "push", "origin", "main"], cwd=PROJECT_DIR, capture_output=True, text=True)
        if result.returncode != 0:
            return jsonify({"error": "Error al hacer push a GitHub", "details": result.stderr}), 500
        
        log_activity("Despliegue WEB", "Actualización forzada a Vercel")
        return jsonify({"ok": True, "message": "Cambios enviados a Vercel con éxito. Estarán visibles en ~1 minuto."})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/admin/log", methods=["GET"])
def admin_get_log():
    """Return the last 100 lines of the activity log."""
    if not check_admin_password():
        return jsonify({"error": "No autorizado"}), 401

    entries = []
    if os.path.exists(LOG_PATH):
        with open(LOG_PATH, "r", encoding="utf-8") as f:
            entries = f.readlines()[-100:]
    return jsonify({"entries": [e.strip() for e in entries if e.strip()]})



@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({
        "status": "ok",
        "contaminants": len(contaminants),
        "templates": {
            "F01655": os.path.exists(F01655_PATH),
            "F00662": os.path.exists(F00662_PATH),
        }
    })


if __name__ == "__main__":
    print("🔬 Gestión Laboratorio — Higiene Industrial API")
    print(f"   Contaminantes: {len(contaminants)}")
    print(f"   F01655: {'✅' if os.path.exists(F01655_PATH) else '❌'}")
    print(f"   F00662: {'✅' if os.path.exists(F00662_PATH) else '❌'}")
    print()
    app.run(host="0.0.0.0", port=5003, debug=True)
